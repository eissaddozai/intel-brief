#!/usr/bin/env python3
"""
export_brief.py — Export a BriefCycle JSON to standalone HTML and Word (.docx).
Usage:
    python scripts/export_brief.py                              # latest brief
    python scripts/export_brief.py briefs/CSE_Intel_Brief_20260308_224500.json
"""

import json
import os
import sys
from pathlib import Path
from datetime import datetime

REPO_ROOT = Path(__file__).resolve().parent.parent
BRIEFS_DIR = REPO_ROOT / "briefs"

# ─── Domain accent colours ──────────────────────────────────────────────────
DOMAIN_COLORS = {
    "d1": {"bg": "#4a1a1a", "accent": "#c0392b", "name": "crimson"},
    "d2": {"bg": "#4a3a1a", "accent": "#d4a017", "name": "amber"},
    "d3": {"bg": "#1a3a2a", "accent": "#27ae60", "name": "green"},
    "d4": {"bg": "#2a1a3a", "accent": "#8e44ad", "name": "purple"},
    "d5": {"bg": "#1a2a3a", "accent": "#4682b4", "name": "steel-blue"},
    "d6": {"bg": "#1a2a2a", "accent": "#2c8c99", "name": "teal"},
}

LEVEL_COLORS = {
    "RED": "#c0392b",
    "AMBER": "#d4a017",
    "GREEN": "#27ae60",
}


def find_latest_brief() -> Path:
    briefs = sorted(BRIEFS_DIR.glob("CSE_Intel_Brief_*.json"), reverse=True)
    if not briefs:
        print("No briefs found.")
        sys.exit(1)
    return briefs[0]


def load_brief(path: Path) -> dict:
    return json.loads(path.read_text("utf-8"))


# ═══════════════════════════════════════════════════════════════════════════════
# HTML EXPORT
# ═══════════════════════════════════════════════════════════════════════════════

def export_html(brief: dict, out_path: Path):
    meta = brief["meta"]
    sh = brief.get("strategicHeader", {})
    ex = brief.get("executive", {})
    domains = brief.get("domains", [])
    wis = brief.get("warningIndicators", [])
    gaps = brief.get("collectionGaps", [])
    cav = brief.get("caveats", {})
    footer = brief.get("footer", {})
    fps = brief.get("flashPoints", [])

    strip_cells_html = ""
    for cell in meta.get("stripCells", []):
        strip_cells_html += f'<div class="strip-cell"><div class="strip-top">{cell["top"]}</div><div class="strip-bot">{cell["bot"]}</div></div>\n'

    flash_html = ""
    for fp in fps:
        dc = DOMAIN_COLORS.get(fp.get("domain", "d1"), DOMAIN_COLORS["d1"])
        flash_html += f'''
        <div class="flash-point" style="border-left: 4px solid {dc["accent"]};">
            <div class="fp-headline">{fp["headline"]}</div>
            <div class="fp-detail">{fp["detail"]}</div>
        </div>'''

    kj_html = ""
    for kj in ex.get("keyJudgments", []):
        dc = DOMAIN_COLORS.get(kj.get("domain", "d1"), DOMAIN_COLORS["d1"])
        kj_html += f'''
        <div class="kj-item" style="border-left: 3px solid {dc["accent"]};">
            <span class="kj-domain">{kj.get("domain","").upper()}</span>
            <span class="kj-conf">[{kj.get("confidence","?")}]</span>
            <span class="kj-text">{kj["text"]}</span>
        </div>'''

    kpi_html = ""
    for kpi in ex.get("kpis", []):
        arrow = {"up": "&#9650;", "down": "&#9660;", "neutral": "&#9654;"}.get(kpi.get("changeDirection", "neutral"), "")
        kpi_html += f'<div class="kpi"><div class="kpi-num">{kpi["number"]}</div><div class="kpi-label">{kpi["label"]}</div><div class="kpi-arrow">{arrow}</div></div>\n'

    domain_sections_html = ""
    for d in domains:
        dc = DOMAIN_COLORS.get(d["id"], DOMAIN_COLORS["d1"])
        paras_html = ""
        for p in d.get("bodyParagraphs", []):
            variant = p.get("subLabelVariant", "observed")
            label_class = "sublabel-observed" if variant == "observed" else "sublabel-assessment"
            cites = ", ".join(
                f'{c.get("source","?")} (T{c.get("tier","?")})'
                for c in p.get("citations", [])
            )
            paras_html += f'''
            <div class="body-para">
                <div class="{label_class}">{p.get("subLabel","")}</div>
                <p>{p.get("text","")}</p>
                <div class="citations">{cites}</div>
            </div>'''

        dissenter_html = ""
        dn = d.get("dissenterNote")
        if dn:
            dissenter_html = f'''
            <div class="dissenter-note">
                <div class="dissenter-label">{dn.get("analyst","ANALYST B")} — DISSENTING VIEW</div>
                <p>{dn.get("text","")}</p>
            </div>'''

        kj_d = d.get("keyJudgment", {})
        domain_sections_html += f'''
        <section class="domain domain--{d["id"]}">
            <div class="domain-header" style="background: {dc["bg"]}; border-bottom: 3px solid {dc["accent"]};">
                <span class="domain-num">{d["num"]}</span>
                <span class="domain-title">{d["title"]}</span>
                <span class="domain-conf">[{d.get("confidence","?")}]</span>
            </div>
            <div class="domain-aq">{d.get("assessmentQuestion","")}</div>
            <div class="domain-kj" style="border-left: 4px solid {dc["accent"]};">
                <strong>KEY JUDGMENT:</strong> {kj_d.get("text","")}
            </div>
            {paras_html}
            {dissenter_html}
            <hr class="section-end" style="border-color: {dc["accent"]};">
        </section>'''

    wi_html = ""
    for wi in wis:
        level = wi.get("level", "AMBER")
        color = LEVEL_COLORS.get(level, "#d4a017")
        wi_html += f'''
        <div class="wi-item">
            <span class="wi-level" style="background:{color};">{level}</span>
            <span class="wi-indicator">{wi.get("indicator","")}</span>
            <p class="wi-assessment">{wi.get("assessment","")}</p>
        </div>'''

    gap_html = ""
    for g in gaps:
        gap_html += f'''
        <div class="gap-item">
            <span class="gap-severity gap-{g.get("severity","minor")}">{g.get("severity","").upper()}</span>
            <span class="gap-domain">[{g.get("domain","")}]</span>
            {g.get("gap","")}
            <p class="gap-sig">{g.get("significance","")}</p>
        </div>'''

    caveat_items = ""
    for c in cav.get("items", []):
        caveat_items += f'<div class="caveat"><strong>{c["label"]}:</strong> {c["text"]}</div>\n'

    html = f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>CSE Intelligence Brief — {meta.get("cycleId","")}</title>
<style>
  @import url('https://fonts.googleapis.com/css2?family=IBM+Plex+Mono:wght@400;600&display=swap');

  :root {{
    --bg: #0a0a12;
    --surface: #12121e;
    --text: #d4d4d4;
    --text-dim: #8a8a9a;
    --accent-red: #c0392b;
    --accent-gold: #d4a017;
    --rule: #2a2a3a;
  }}

  * {{ margin: 0; padding: 0; box-sizing: border-box; }}

  body {{
    font-family: Georgia, 'Times New Roman', serif;
    background: var(--bg);
    color: var(--text);
    line-height: 1.6;
    max-width: 900px;
    margin: 0 auto;
    padding: 2rem 1.5rem;
  }}

  /* ─── Cover ─── */
  .cover {{
    background: var(--surface);
    border: 1px solid var(--rule);
    border-top: 4px solid var(--accent-red);
    padding: 2rem;
    margin-bottom: 2rem;
    text-align: center;
  }}
  .cover h1 {{ font-family: Palatino, Georgia, serif; font-size: 1.8rem; color: #e8e8e8; letter-spacing: 0.15em; }}
  .cover .subtitle {{ font-size: 0.95rem; color: var(--accent-red); margin: 0.5rem 0; font-weight: 600; }}
  .cover .meta-line {{ font-size: 0.75rem; color: var(--text-dim); font-style: italic; }}
  .cover .context {{ font-size: 0.8rem; color: var(--text-dim); margin-top: 1rem; text-align: left; }}

  /* ─── Strip cells ─── */
  .strip {{ display: flex; gap: 0.5rem; margin: 1.5rem 0; flex-wrap: wrap; justify-content: center; }}
  .strip-cell {{ background: var(--surface); border: 1px solid var(--rule); padding: 0.6rem 1rem; text-align: center; min-width: 120px; }}
  .strip-top {{ font-family: 'IBM Plex Mono', monospace; font-weight: 600; font-size: 1rem; color: #e8e8e8; }}
  .strip-bot {{ font-family: 'IBM Plex Mono', monospace; font-size: 0.65rem; color: var(--text-dim); letter-spacing: 0.1em; }}

  /* ─── Classification ─── */
  .classification {{
    font-family: 'IBM Plex Mono', monospace;
    font-size: 0.7rem;
    color: var(--accent-gold);
    text-align: center;
    letter-spacing: 0.15em;
    margin: 1rem 0;
    border: 1px solid var(--accent-gold);
    padding: 0.3rem;
  }}

  /* ─── Strategic header ─── */
  .strategic {{ background: var(--surface); border-left: 4px solid var(--accent-red); padding: 1.2rem 1.5rem; margin: 1.5rem 0; }}
  .strategic h2 {{ font-family: Palatino, Georgia, serif; font-size: 0.75rem; color: var(--accent-red); letter-spacing: 0.15em; margin-bottom: 0.5rem; }}
  .strategic .headline {{ font-size: 1rem; color: #e8e8e8; font-weight: 600; }}
  .strategic .trajectory {{ font-size: 0.85rem; color: var(--text); margin-top: 0.8rem; }}

  /* ─── Flash points ─── */
  .flash-section h2 {{ font-family: 'IBM Plex Mono', monospace; font-size: 0.75rem; color: var(--accent-red); letter-spacing: 0.15em; margin-bottom: 1rem; }}
  .flash-point {{ background: var(--surface); padding: 1rem 1.2rem; margin-bottom: 0.8rem; }}
  .fp-headline {{ font-weight: 600; color: #e8e8e8; font-size: 0.95rem; margin-bottom: 0.4rem; }}
  .fp-detail {{ font-size: 0.85rem; color: var(--text); }}

  /* ─── Executive ─── */
  .executive {{ margin: 2rem 0; }}
  .executive h2 {{ font-family: 'IBM Plex Mono', monospace; font-size: 0.75rem; color: var(--accent-gold); letter-spacing: 0.15em; margin-bottom: 0.8rem; }}
  .bluf {{ font-size: 0.9rem; color: var(--text); margin-bottom: 1.5rem; line-height: 1.7; }}
  .kj-item {{ padding: 0.6rem 1rem; margin-bottom: 0.5rem; background: var(--surface); }}
  .kj-domain {{ font-family: 'IBM Plex Mono', monospace; font-size: 0.7rem; color: var(--text-dim); }}
  .kj-conf {{ font-family: 'IBM Plex Mono', monospace; font-size: 0.65rem; color: var(--accent-gold); }}
  .kj-text {{ font-size: 0.85rem; }}
  .kpi-strip {{ display: flex; gap: 0.5rem; margin: 1.5rem 0; flex-wrap: wrap; }}
  .kpi {{ background: var(--surface); border: 1px solid var(--rule); padding: 0.6rem 1rem; text-align: center; min-width: 140px; flex: 1; }}
  .kpi-num {{ font-family: 'IBM Plex Mono', monospace; font-weight: 600; font-size: 1.1rem; color: #e8e8e8; }}
  .kpi-label {{ font-size: 0.6rem; color: var(--text-dim); letter-spacing: 0.1em; }}
  .kpi-arrow {{ font-size: 0.7rem; color: var(--accent-red); }}

  /* ─── Domain sections ─── */
  .domain {{ margin: 2rem 0; }}
  .domain-header {{ padding: 0.8rem 1.2rem; display: flex; align-items: baseline; gap: 0.8rem; }}
  .domain-num {{ font-family: 'IBM Plex Mono', monospace; font-size: 1.5rem; font-weight: 600; color: #e8e8e8; opacity: 0.5; }}
  .domain-title {{ font-family: Palatino, Georgia, serif; font-size: 1.1rem; color: #e8e8e8; letter-spacing: 0.1em; }}
  .domain-conf {{ font-family: 'IBM Plex Mono', monospace; font-size: 0.7rem; color: var(--accent-gold); }}
  .domain-aq {{ font-size: 0.8rem; color: var(--text-dim); font-style: italic; padding: 0.5rem 1.2rem; background: rgba(255,255,255,0.02); }}
  .domain-kj {{ background: var(--surface); padding: 1rem 1.2rem; margin: 0.8rem 0; font-size: 0.9rem; }}
  .body-para {{ padding: 0.8rem 1.2rem; margin-bottom: 0.5rem; }}
  .body-para p {{ font-size: 0.85rem; line-height: 1.7; }}
  .sublabel-observed {{ font-family: 'IBM Plex Mono', monospace; font-size: 0.65rem; color: var(--text-dim); letter-spacing: 0.12em; margin-bottom: 0.3rem; }}
  .sublabel-assessment {{ font-family: 'IBM Plex Mono', monospace; font-size: 0.65rem; color: var(--accent-gold); letter-spacing: 0.12em; margin-bottom: 0.3rem; opacity: 0.7; }}
  .citations {{ font-size: 0.7rem; color: var(--text-dim); font-style: italic; margin-top: 0.3rem; }}
  .dissenter-note {{ background: #1a1a2e; border: 1px dashed var(--accent-gold); padding: 0.8rem 1rem; margin: 0.8rem 0; font-size: 0.8rem; }}
  .dissenter-label {{ font-family: 'IBM Plex Mono', monospace; font-size: 0.65rem; color: var(--accent-gold); margin-bottom: 0.3rem; }}
  .section-end {{ border: none; border-top: 2px solid; margin: 1rem 0; }}

  /* ─── Warning indicators ─── */
  .wi-section h2 {{ font-family: 'IBM Plex Mono', monospace; font-size: 0.75rem; color: var(--accent-red); letter-spacing: 0.15em; margin-bottom: 1rem; }}
  .wi-item {{ padding: 0.6rem 0; border-bottom: 1px solid var(--rule); }}
  .wi-level {{ font-family: 'IBM Plex Mono', monospace; font-size: 0.65rem; color: #fff; padding: 0.15rem 0.5rem; border-radius: 2px; margin-right: 0.5rem; }}
  .wi-indicator {{ font-weight: 600; font-size: 0.85rem; }}
  .wi-assessment {{ font-size: 0.8rem; color: var(--text-dim); margin-top: 0.2rem; padding-left: 1rem; }}

  /* ─── Collection gaps ─── */
  .gaps-section h2 {{ font-family: 'IBM Plex Mono', monospace; font-size: 0.75rem; color: var(--text-dim); letter-spacing: 0.15em; margin-bottom: 1rem; }}
  .gap-item {{ padding: 0.5rem 0; border-bottom: 1px solid var(--rule); font-size: 0.85rem; }}
  .gap-severity {{ font-family: 'IBM Plex Mono', monospace; font-size: 0.6rem; padding: 0.1rem 0.4rem; border-radius: 2px; margin-right: 0.3rem; }}
  .gap-critical {{ background: var(--accent-red); color: #fff; }}
  .gap-significant {{ background: var(--accent-gold); color: #1a1a2e; }}
  .gap-minor {{ background: var(--rule); color: var(--text-dim); }}
  .gap-domain {{ font-family: 'IBM Plex Mono', monospace; font-size: 0.7rem; color: var(--text-dim); }}
  .gap-sig {{ font-size: 0.8rem; color: var(--text-dim); padding-left: 1rem; }}

  /* ─── Caveats ─── */
  .caveats {{ background: var(--surface); border: 1px solid var(--rule); padding: 1.5rem; margin: 2rem 0; }}
  .caveats h2 {{ font-family: 'IBM Plex Mono', monospace; font-size: 0.75rem; color: var(--text-dim); letter-spacing: 0.15em; margin-bottom: 1rem; }}
  .caveat {{ font-size: 0.8rem; margin-bottom: 0.5rem; }}
  .conf-assessment {{ font-size: 0.8rem; color: var(--text-dim); margin-top: 0.8rem; font-style: italic; }}

  /* ─── Footer ─── */
  .footer {{
    text-align: center;
    font-family: 'IBM Plex Mono', monospace;
    font-size: 0.65rem;
    color: var(--text-dim);
    border-top: 2px solid var(--rule);
    padding-top: 1rem;
    margin-top: 2rem;
    letter-spacing: 0.1em;
  }}

  @media print {{
    body {{ background: #fff; color: #1a1a1a; max-width: 100%; }}
    .cover {{ border-color: #1a1a1a; }}
    .strip-cell, .kpi, .domain-kj, .flash-point, .kj-item {{ border-color: #ccc; }}
  }}
</style>
</head>
<body>

<div class="classification">{meta.get("classification","")} // TLP:{meta.get("tlp","")}</div>

<div class="cover">
    <h1>CSE INTELLIGENCE BRIEF</h1>
    <div class="subtitle">{meta.get("subtitle","")}</div>
    <div class="meta-line">{meta.get("timestamp","")[:10]} &middot; Cycle {meta.get("cycleNum","")} &middot; {meta.get("region","")}</div>
    <div class="meta-line">{meta.get("analystUnit","")}</div>
    <div class="context">{meta.get("contextNote","")}</div>
</div>

<div class="strip">
{strip_cells_html}
</div>

<div class="strategic">
    <h2>HEADLINE JUDGMENT</h2>
    <div class="headline">{sh.get("headlineJudgment","")}</div>
    <div class="trajectory">{sh.get("trajectoryRationale","")}</div>
</div>

<div class="flash-section">
    <h2>FLASH POINTS</h2>
    {flash_html}
</div>

<div class="executive">
    <h2>EXECUTIVE SUMMARY</h2>
    <div class="bluf">{ex.get("bluf","")}</div>
    <h2>KEY JUDGMENTS</h2>
    {kj_html}
    <div class="kpi-strip">
    {kpi_html}
    </div>
</div>

{domain_sections_html}

<div class="wi-section">
    <h2>WARNING INDICATORS</h2>
    {wi_html}
</div>

<div class="gaps-section">
    <h2>COLLECTION GAPS</h2>
    {gap_html}
</div>

<div class="caveats">
    <h2>CAVEATS &amp; SOURCE QUALITY</h2>
    <div class="caveat"><strong>Cycle:</strong> {cav.get("cycleRef","")}</div>
    {caveat_items}
    <div class="conf-assessment">{cav.get("confidenceAssessment","")}</div>
    <div class="caveat" style="margin-top:0.8rem;"><strong>Source Quality:</strong> {cav.get("sourceQuality","")}</div>
    <div class="caveat"><strong>Handling:</strong> {cav.get("handling","")}</div>
</div>

<div class="footer">
    {footer.get("classification","")} &middot; {footer.get("id","")}<br>
    Sources: {footer.get("sources","")}<br>
    {footer.get("handling","")}
</div>

</body>
</html>"""

    out_path.write_text(html, encoding="utf-8")
    print(f"HTML exported: {out_path}")


# ═══════════════════════════════════════════════════════════════════════════════
# WORD EXPORT
# ═══════════════════════════════════════════════════════════════════════════════

def export_docx(brief: dict, out_path: Path):
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    def set_cell_bg(cell, hex_color: str):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color)
        tcPr.append(shd)

    def add_heading_styled(doc, text, level, color_hex="1A1A2E"):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt([0, 20, 16, 13, 11][level])
        r, g, b = tuple(int(color_hex[i : i + 2], 16) for i in (0, 2, 4))
        run.font.color.rgb = RGBColor(r, g, b)
        return p

    def add_body_text(doc, text, italic=False, bold=False, size=10):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.italic = italic
        run.bold = bold
        run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
        return p

    def add_hr(doc):
        p = doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        pb = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "C0392B")
        pb.append(bottom)
        pPr.append(pb)

    meta = brief["meta"]
    sh = brief.get("strategicHeader", {})
    ex = brief.get("executive", {})
    domains = brief.get("domains", [])
    wis = brief.get("warningIndicators", [])
    gaps = brief.get("collectionGaps", [])
    cav = brief.get("caveats", {})
    footer = brief.get("footer", {})
    fps = brief.get("flashPoints", [])

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    # ─── COVER ───
    cover = doc.add_table(rows=1, cols=1)
    cover.style = "Table Grid"
    cell = cover.rows[0].cells[0]
    set_cell_bg(cell, "1A1A2E")
    cp = cell.paragraphs[0]
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.space_before = Pt(12)
    r1 = cp.add_run("\nCSE INTELLIGENCE BRIEF\n")
    r1.bold = True
    r1.font.size = Pt(22)
    r1.font.color.rgb = RGBColor(0xE8, 0xE8, 0xE8)

    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(meta.get("subtitle", ""))
    r2.bold = True
    r2.font.size = Pt(14)
    r2.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)

    p3 = cell.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ts = meta.get("timestamp", "")[:10]
    r3 = p3.add_run(
        f'{ts}  |  Cycle {meta.get("cycleNum","")}  |  {meta.get("classification","")} // TLP:{meta.get("tlp","")}'
    )
    r3.font.size = Pt(9)
    r3.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
    r3.italic = True

    p4 = cell.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p4.paragraph_format.space_after = Pt(12)
    r4 = p4.add_run(f'Sources: {footer.get("sources","")}')
    r4.font.size = Pt(7.5)
    r4.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_paragraph()

    # ─── STRIP CELLS ───
    cells_data = meta.get("stripCells", [])
    if cells_data:
        strip_table = doc.add_table(rows=2, cols=len(cells_data))
        strip_table.style = "Table Grid"
        for i, sc in enumerate(cells_data):
            set_cell_bg(strip_table.rows[0].cells[i], "2C3E50")
            rr = strip_table.rows[0].cells[i].paragraphs[0].add_run(sc["top"])
            rr.bold = True
            rr.font.size = Pt(11)
            rr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            strip_table.rows[0].cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            rr2 = strip_table.rows[1].cells[i].paragraphs[0].add_run(sc["bot"])
            rr2.font.size = Pt(7)
            rr2.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)
            strip_table.rows[1].cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()

    # ─── STRATEGIC HEADER ───
    add_heading_styled(doc, "HEADLINE JUDGMENT", 2, "C0392B")
    add_hr(doc)
    add_body_text(doc, sh.get("headlineJudgment", ""), bold=True)
    add_body_text(doc, sh.get("trajectoryRationale", ""), italic=True, size=9)
    doc.add_paragraph()

    # ─── FLASH POINTS ───
    if fps:
        add_heading_styled(doc, "FLASH POINTS", 2, "C0392B")
        add_hr(doc)
        for fp in fps:
            p = doc.add_paragraph()
            r = p.add_run(f'[{fp.get("domain","").upper()}] {fp["headline"]}')
            r.bold = True
            r.font.size = Pt(11)
            r.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
            add_body_text(doc, fp.get("detail", ""), size=9)
        doc.add_paragraph()

    # ─── EXECUTIVE ───
    add_heading_styled(doc, "EXECUTIVE SUMMARY", 2, "1A1A2E")
    add_hr(doc)
    add_body_text(doc, ex.get("bluf", ""))

    add_heading_styled(doc, "Key Judgments", 3, "1A1A2E")
    for kj in ex.get("keyJudgments", []):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        r1 = p.add_run(f'[{kj.get("domain","").upper()}] ')
        r1.bold = True
        r1.font.size = Pt(9)
        r1.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)
        r2 = p.add_run(kj["text"])
        r2.font.size = Pt(10)
        r2.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
    doc.add_paragraph()

    # ─── DOMAIN SECTIONS ───
    DOMAIN_DOCX_COLORS = {
        "d1": "C0392B", "d2": "D4A017", "d3": "27AE60",
        "d4": "8E44AD", "d5": "4682B4", "d6": "2C8C99",
    }

    for d in domains:
        color = DOMAIN_DOCX_COLORS.get(d["id"], "1A1A2E")
        table = doc.add_table(rows=1, cols=1)
        table.style = "Table Grid"
        cell = table.rows[0].cells[0]
        set_cell_bg(cell, color)
        p = cell.paragraphs[0]
        r = p.add_run(f'  {d["num"]} — {d["title"]}')
        r.bold = True
        r.font.size = Pt(13)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        doc.add_paragraph()

        # AQ
        add_body_text(doc, d.get("assessmentQuestion", ""), italic=True, size=9)

        # KJ
        kj_d = d.get("keyJudgment", {})
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        r = p.add_run("KEY JUDGMENT: ")
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
        r2 = p.add_run(kj_d.get("text", ""))
        r2.font.size = Pt(10)
        r2.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)

        # Body
        for bp in d.get("bodyParagraphs", []):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            r_label = p.add_run(f'[{bp.get("subLabel","")}] ')
            r_label.bold = True
            r_label.font.size = Pt(8)
            r_label.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)
            r_text = p.add_run(bp.get("text", ""))
            r_text.font.size = Pt(9)
            r_text.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)

            cites = [c.get("source", "") for c in bp.get("citations", [])]
            if cites:
                p2 = doc.add_paragraph()
                r3 = p2.add_run("Sources: " + " | ".join(cites))
                r3.italic = True
                r3.font.size = Pt(7)
                r3.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

        # Dissenter
        dn = d.get("dissenterNote")
        if dn:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            r = p.add_run(f'{dn.get("analyst","ANALYST B")} — DISSENTING VIEW: ')
            r.bold = True
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(0xD4, 0xA0, 0x17)
            r2 = p.add_run(dn.get("text", ""))
            r2.font.size = Pt(9)
            r2.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)

        doc.add_paragraph()

    # ─── WARNING INDICATORS ───
    add_heading_styled(doc, "WARNING INDICATORS", 2, "C0392B")
    add_hr(doc)
    for wi in wis:
        p = doc.add_paragraph()
        r = p.add_run(f'[{wi.get("level","")}] {wi.get("indicator","")}')
        r.bold = True
        r.font.size = Pt(10)
        color = {"RED": 0xC0392B, "AMBER": 0xD4A017, "GREEN": 0x27AE60}.get(
            wi.get("level", "AMBER"), 0xD4A017
        )
        r.font.color.rgb = RGBColor((color >> 16) & 0xFF, (color >> 8) & 0xFF, color & 0xFF)
        add_body_text(doc, wi.get("assessment", ""), size=9)
    doc.add_paragraph()

    # ─── COLLECTION GAPS ───
    add_heading_styled(doc, "COLLECTION GAPS", 2, "1A1A2E")
    for g in gaps:
        p = doc.add_paragraph()
        r = p.add_run(f'[{g.get("severity","").upper()}] [{g.get("domain","")}] ')
        r.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)
        r2 = p.add_run(g.get("gap", ""))
        r2.font.size = Pt(9)
        add_body_text(doc, g.get("significance", ""), italic=True, size=8)
    doc.add_paragraph()

    # ─── CAVEATS ───
    add_heading_styled(doc, "CAVEATS & SOURCE QUALITY", 2, "1A1A2E")
    add_hr(doc)
    for c in cav.get("items", []):
        p = doc.add_paragraph()
        r = p.add_run(f'{c["label"]}: ')
        r.bold = True
        r.font.size = Pt(9)
        r2 = p.add_run(c["text"])
        r2.font.size = Pt(9)
        r2.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
    add_body_text(doc, cav.get("confidenceAssessment", ""), italic=True, size=9)
    add_body_text(doc, f'Source Quality: {cav.get("sourceQuality","")}', size=8)
    doc.add_paragraph()

    # ─── FOOTER ───
    add_hr(doc)
    fp = doc.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = fp.add_run(
        f'{footer.get("classification","")}\n'
        f'{footer.get("handling","")}'
    )
    r.italic = True
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)

    doc.save(str(out_path))
    print(f"Word exported: {out_path}")


# ═══════════════════════════════════════════════════════════════════════════════
# MAIN
# ═══════════════════════════════════════════════════════════════════════════════

def main():
    if len(sys.argv) > 1:
        brief_path = Path(sys.argv[1])
        if not brief_path.is_absolute():
            brief_path = REPO_ROOT / brief_path
    else:
        brief_path = find_latest_brief()

    print(f"Reading: {brief_path.name}")
    brief = load_brief(brief_path)
    stem = brief_path.stem

    html_out = BRIEFS_DIR / f"{stem}.html"
    docx_out = BRIEFS_DIR / f"{stem}.docx"

    export_html(brief, html_out)
    export_docx(brief, docx_out)


if __name__ == "__main__":
    main()
