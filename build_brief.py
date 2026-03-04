"""
build_brief.py — Compile today's intelligence brief into a Word document.
Run: python3 build_brief.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime, os

# ─── helpers ─────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    """Set table cell background color."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def add_horizontal_rule(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pb = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'C0392B')
    pb.append(bottom)
    pPr.append(pb)
    return p


def add_heading(doc, text: str, level: int, color_hex: str = '1A1A2E', bold: bool = True):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt([0, 20, 16, 13, 11][level])
    r, g, b = tuple(int(color_hex[i:i+2], 16) for i in (0, 2, 4))
    run.font.color.rgb = RGBColor(r, g, b)
    return p


def add_body(doc, text: str, italic: bool = False, bold: bool = False,
             indent: float = 0, space_after: int = 4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    if indent:
        p.paragraph_format.left_indent = Inches(indent)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.italic = italic
    run.bold = bold
    run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
    return p


def add_bullet(doc, text: str, indent: float = 0.2, bold_prefix: str = ''):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Inches(indent)
    if bold_prefix:
        run_b = p.add_run(bold_prefix)
        run_b.bold = True
        run_b.font.size = Pt(10)
        run = p.add_run(text)
    else:
        run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
    return p


def add_source_note(doc, sources: list[str]):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run('Sources: ' + ' | '.join(sources))
    run.italic = True
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)
    return p


def add_domain_header(doc, code: str, title: str, color_hex: str):
    """Adds a prominent domain section header as a shaded table cell."""
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.rows[0].cells[0]
    set_cell_bg(cell, color_hex)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.1)
    run = p.add_run(f'  {code} — {title}')
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    doc.add_paragraph()
    return table


def add_callout(doc, label: str, text: str, color_hex: str = 'F8C471'):
    """Inline callout box (single-cell table)."""
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.rows[0].cells[0]
    set_cell_bg(cell, color_hex)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.1)
    r_label = p.add_run(f'{label}  ')
    r_label.bold = True
    r_label.font.size = Pt(10)
    r_label.font.color.rgb = RGBColor(0x78, 0x28, 0x00)
    r_text = p.add_run(text)
    r_text.font.size = Pt(10)
    r_text.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
    doc.add_paragraph()
    return table


# ═══════════════════════════════════════════════════════════════════════════════
# DOCUMENT
# ═══════════════════════════════════════════════════════════════════════════════

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# Default paragraph style
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10)


# ─── COVER BLOCK ──────────────────────────────────────────────────────────────

cover = doc.add_table(rows=1, cols=1)
cover.style = 'Table Grid'
cover_cell = cover.rows[0].cells[0]
set_cell_bg(cover_cell, '1A1A2E')

cp = cover_cell.paragraphs[0]
cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
cp.paragraph_format.space_before = Pt(12)
cp.paragraph_format.space_after = Pt(2)

r1 = cp.add_run('\nCSE INTELLIGENCE BRIEF\n')
r1.bold = True; r1.font.size = Pt(22); r1.font.color.rgb = RGBColor(0xE8, 0xE8, 0xE8)

p2 = cover_cell.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run('IRAN CONFLICT SPACE — MORNING EDITION')
r2.bold = True; r2.font.size = Pt(14); r2.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)

p3 = cover_cell.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p3.add_run('4 MARCH 2026  ·  DAY 5 OF OPERATIONS  ·  UNCLASSIFIED // FOR OFFICIAL USE ONLY')
r3.font.size = Pt(9); r3.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA); r3.font.italic = True

p4 = cover_cell.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
p4.paragraph_format.space_after = Pt(12)
r4 = p4.add_run('Sources consulted: AP · Reuters · CTP-ISW · IAEA · CENTCOM · CNN · BBC · CNBC · Al Jazeera · Times of Israel · Jerusalem Post · NYT · The Guardian · Middle East Eye · DW · France 24 · Iran International · Long War Journal · NetBlocks · CISA · Unit 42 · SC Media · gCaptain · BIMCO · Kpler · Wood Mackenzie')
r4.font.size = Pt(7.5); r4.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_paragraph()


# ─── KEY FACTS TABLE ──────────────────────────────────────────────────────────

add_heading(doc, '◾ SITUATION AT A GLANCE  —  04 MAR 2026  06:00 UTC', 2, '1A1A2E')
add_horizontal_rule(doc)
doc.add_paragraph()

facts_table = doc.add_table(rows=7, cols=2)
facts_table.style = 'Table Grid'
facts_table.autofit = False
facts_table.columns[0].width = Inches(2.2)
facts_table.columns[1].width = Inches(4.4)

rows_data = [
    ('Conflict Day', 'Day 5 — Operation ROARING LION (Israel) / EPIC FURY (US) — Ongoing'),
    ('Iranian KIA (est.)', '1,045+ confirmed dead (Iranian Red Crescent). Thousands IRGC WIA.'),
    ('US KIA', '6 U.S. service members killed; 18 wounded (CENTCOM, 04 Mar)'),
    ('Missiles Fired (Iran)', '500+ ballistic missiles + 2,000+ drones since Feb 28 (CENTCOM Adm. Cooper)'),
    ('Brent Crude', '$83.83/bbl (04 Mar 10:06 UTC) — up ~25% from pre-war level'),
    ('Hormuz Status', 'DE FACTO CLOSED — insurance withdrawal; ~0% commercial transit'),
    ('Iran Internet', '~1% connectivity (NetBlocks, 04 Mar) — regime-imposed total blackout'),
]

header_row = facts_table.rows[0]
set_cell_bg(header_row.cells[0], '2C3E50')
set_cell_bg(header_row.cells[1], '2C3E50')
for i, cell in enumerate(header_row.cells):
    p = cell.paragraphs[0]
    r = p.add_run(['INDICATOR', 'CURRENT STATUS'][i])
    r.bold = True; r.font.size = Pt(9); r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

for row_idx, (label, value) in enumerate(rows_data):
    row = facts_table.rows[row_idx]
    set_cell_bg(row.cells[0], 'ECF0F1')
    p0 = row.cells[0].paragraphs[0]
    r0 = p0.add_run(label)
    r0.bold = True; r0.font.size = Pt(9); r0.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
    p1 = row.cells[1].paragraphs[0]
    r1 = p1.add_run(value)
    r1.font.size = Pt(9); r1.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)

doc.add_paragraph()


# ─── EXECUTIVE SUMMARY ────────────────────────────────────────────────────────

add_heading(doc, '◾ EXECUTIVE SUMMARY', 2, '1A1A2E')
add_horizontal_rule(doc)

add_body(doc,
    'The US–Israeli joint campaign against Iran entered Day 5 on 4 March 2026, '
    'with air operations continuing around the clock. The assassination of Supreme Leader '
    'Ali Khamenei on 1 March has decapitated Iranian state leadership; a provisional '
    'Leadership Council now holds executive authority while the Assembly of Experts — '
    'itself struck by Israeli airstrikes during an emergency succession session — attempts '
    'to convene. Iran has responded with "Operation True Promise IV," firing more than '
    '500 ballistic missiles and 2,000+ drones at Israel, US bases across the Gulf, '
    'and civilian infrastructure in the UAE, Kuwait, Qatar, Bahrain, and Oman.')

add_body(doc,
    'Iran\'s effective closure of the Strait of Hormuz — achieved not by physical naval '
    'blockade but by IRGC threats and the resultant collapse of marine war-risk insurance — '
    'has removed ~20% of global seaborne oil from circulation. Brent crude stands at '
    '$83.83/bbl as of morning 4 March, with Wood Mackenzie warning of $150/bbl if closure '
    'persists beyond two weeks. Hezbollah has re-entered the conflict from Lebanon, '
    'conducting missile and drone strikes against Israeli northern airbases. Iran-backed '
    'Iraqi militias are escalating attacks on US bases in Erbil, Baghdad, and across '
    'Kuwait and Qatar.')

add_body(doc,
    'The IAEA has confirmed it cannot verify Iran\'s nuclear enrichment status, '
    'uranium stockpile size, or current location of fissile material. Iran\'s internet '
    'connectivity stands at approximately 1% of normal per NetBlocks, indicating a regime-'
    'imposed blackout consistent with wartime communications control. Iranian IRGC-affiliated '
    'cyber actors are conducting elevated reconnaissance and disruption operations against '
    'US critical infrastructure. The US holds the UN Security Council presidency this '
    'month; no ceasefire resolution has been tabled.')

doc.add_paragraph()

# ─── D1 BATTLESPACE ───────────────────────────────────────────────────────────

add_domain_header(doc, 'D1', 'BATTLESPACE & MILITARY OPERATIONS', 'C0392B')

add_heading(doc, '1.1  US-ISRAELI STRIKE CAMPAIGN — OVERVIEW', 3, 'C0392B')

add_body(doc,
    'The joint campaign, codenamed Operation ROARING LION (Israel) and EPIC FURY (United States), '
    'was launched on 28 February 2026 from two carrier strike groups — USS Abraham Lincoln '
    '(Arabian Sea) and USS Gerald R. Ford (eastern Mediterranean) — supported by land-based '
    'F-22s deployed to Ovda Airbase, Israel, and F-15Es in Jordan. B-2 stealth bombers '
    'conducted deep-penetrating strikes on hardened ballistic missile facilities.')

add_bullet(doc, 'IAF has conducted 1,600+ sorties into Iranian territory since D1. Over 4,000 munitions dropped — equal to the entire June 2025 twelve-day war.', bold_prefix='Air Operations: ')
add_bullet(doc, 'CENTCOM Adm. Cooper confirmed ~2,000 US strikes on Iranian targets with 2,000+ munitions; 17 Iranian naval vessels destroyed, including the most capable operational Iranian submarine.', bold_prefix='Naval: ')
add_bullet(doc, '50,000+ US troops and 200+ aircraft deployed across the theater (as reported).', bold_prefix='Force Size: ')
add_bullet(doc, 'Joint Chiefs Chair Gen. Dan Caine: "local air superiority has been established" over western Iran and Tehran as of Day 3.', bold_prefix='Air Superiority: ')
add_bullet(doc, 'USS Gerald R. Ford positioned off Israeli coast. USS Abraham Lincoln in Arabian Sea. 16 surface warships in theater — largest US naval presence in region since OIF 2003.', bold_prefix='Naval Posture: ')

add_source_note(doc, ['CTP-ISW Special Reports Mar 1–4', 'CENTCOM', 'Times of Israel', 'Military Times', 'CNN'])

add_heading(doc, '1.2  KEY TARGETS STRUCK (DAYS 1–5)', 3, 'C0392B')

add_bullet(doc, 'Khamenei compound, Tehran — Supreme Leader Ali Khamenei killed D2 (confirmed by Iranian state media 1 Mar). Wife also killed of injuries.')
add_bullet(doc, 'IRGC Malek-Ashtar defense research complex, Tehran — destroyed (footage confirmed D3).')
add_bullet(doc, 'Islamic Republic of Iran Broadcasting (IRIB) HQ, Tehran — struck D4.')
add_bullet(doc, 'Iran\'s parliament building — targeted D4.')
add_bullet(doc, 'Assembly of Experts — struck during emergency meeting to elect successor supreme leader.')
add_bullet(doc, 'IRGC Quds Force HQ, Tehran — struck D3.')
add_bullet(doc, '10 Intelligence Ministry command centers — struck D4.')
add_bullet(doc, '2nd Artesh Air Force Tactical Airbase, Tabriz — satellite imagery confirms 11 runway craters; airbase inoperable.')
add_bullet(doc, 'Natanz enrichment facility — significant damage to entrance structures; IAEA confirms inaccessible. Enrichment status unverifiable.')
add_bullet(doc, 'Nuclear weaponization research facilities — IDF confirmed strikes D3-4.')
add_bullet(doc, 'US B-2 bombers: struck ballistic missile production facilities with 2,000 lb GBU-28 penetrating munitions.')

add_source_note(doc, ['CTP-ISW', 'CENTCOM', 'Times of Israel', 'IAEA', 'Al Jazeera', 'CNN'])

add_heading(doc, '1.3  IRANIAN BALLISTIC MISSILE STOCKPILE & ATTRITION', 3, 'C0392B')

add_body(doc,
    'Pre-war Israeli intelligence assessed Iran held approximately 2,500 ballistic missiles '
    'following partial replenishment from the June 2025 twelve-day war, down from a '
    'pre-2025 peak of ~3,000. Iran was accelerating production toward an 8,000-unit '
    'target before hostilities.')

add_bullet(doc, '500+ ballistic missiles and 2,000+ drones fired since D1 (CENTCOM Adm. Cooper, D4).')
add_bullet(doc, '~300 Iranian missile launchers destroyed per IDF assessment (D4).')
add_bullet(doc, 'Iranian missile attack frequency against Israel has "decreased sharply" — consistent with launcher destruction degrading launch tempo.', bold_prefix='Key Signal: ')
add_bullet(doc, 'Times of Israel liveblog (04 Mar): "Within days, Iran may run out of ballistic missiles; US may have to start conserving interceptors."')
add_bullet(doc, 'US SM-3 and Tomahawk stocks at risk of a "Winchester" scenario — depletion may force diversion from Pacific, degrading China deterrence.', bold_prefix='US Risk: ')
add_bullet(doc, 'Iran retains Sejil, Khorramshahr, Kheibar class missiles (2,000–2,500 km range). Mobile TEL launchers remain hard to neutralize; emerge-launch-withdraw within minutes.', bold_prefix='Survivable Inventory: ')

add_callout(doc, '⚠ ATTRITION SIGNAL:',
    'The rate of Iranian launches is declining. If this is launcher-driven rather than command-decision-driven, '
    'Iran\'s offensive missile capacity may reach a tipping point within 3–5 days at current burn rates. '
    'This could force Iran to negotiate or shift entirely to asymmetric/proxy operations.',
    'FADBD1')

add_source_note(doc, ['Times of Israel', 'CTP-ISW Mar 3–4', 'CNN Politics missile math', 'ABC News', 'Interesting Engineering'])

add_heading(doc, '1.4  IRAN COUNTERSTRIKES — "OPERATION TRUE PROMISE IV"', 3, 'C0392B')

add_body(doc, 'Iran launched "Operation True Promise IV" — waves of ballistic missiles and Shahed-series drones across the entire Gulf AOR:')

add_bullet(doc, 'Israel: Multiple salvos; Iron Dome, Arrow-3, David\'s Sling, and THAAD batteries engaged. Northern targets: Ramat David Airbase, Meron monitoring base, Camp Yitzhak (Hezbollah-coordinated).')
add_bullet(doc, 'Qatar (Al Udeid AB): 104 ballistic missiles + 39 drones detected; 101 missiles and 24 drones intercepted. 2 Iranian SU-24 aircraft shot down by Qatar. IRGC cell arrested inside Qatar.', bold_prefix='Qatar: ')
add_bullet(doc, 'Kuwait: 178 BMs + 384 drones intercepted. Iranian drone struck Camp Buehring. 6 US service members killed at Kuwait base — highest single-location US fatality event.', bold_prefix='Kuwait: ')
add_bullet(doc, 'UAE: 174 BMs (13 not intercepted; fell into sea), 8 cruise missiles (all intercepted), 689 UAVs (44 caused impact). Dubai International Airport struck — 4 staff injured. Fairmont Palm Hotel, Palm Jumeirah: Shahed drone strike, explosion and fire. Burj Al Arab: impact debris damage. Jebel Ali Port: debris fires. AWS/Amazon data center mec1-az2 struck, fire, power shutdown.', bold_prefix='UAE: ')
add_bullet(doc, 'Bahrain (Fifth Fleet HQ): targeted. USS vessels pre-departed port prior to strikes.', bold_prefix='Bahrain: ')
add_bullet(doc, 'Saudi Arabia: US Embassy Riyadh struck by 2 Iranian drones (March 2). 4 PMU personnel killed in Diyala, Iraq by US/Israeli strikes on PMF HQ.', bold_prefix='Saudi Arabia / Iraq: ')
add_bullet(doc, 'Oman: Iranian drone struck Port of Salalah; Iranian drones struck fuel tanker at Duqm Port.', bold_prefix='Oman: ')
add_bullet(doc, 'Cyprus: UK\'s Akrotiri and Dhekelia military bases struck.', bold_prefix='Cyprus: ')

add_source_note(doc, ['Al Jazeera', 'CNN', 'CENTCOM', 'Euronews', 'Breaking Defense', 'CNBC', 'Wikipedia'])

add_heading(doc, '1.5  IRAN NAVAL ACTION — IRIS DENA SUNK', 3, 'C0392B')

add_body(doc,
    'On 4 March 2026, CTP-ISW reported that the Iranian Navy frigate IRIS Dena was '
    'reportedly sunk in the Indian Ocean by US Navy submarines, approximately 40 nautical '
    'miles south of Galle, Sri Lanka. The vessel was transiting home following participation '
    'in Exercise MILAN at Visakhapatnam, India. Sri Lanka Navy and Air Force conducted '
    'SAR operations; vessel sank before rescue forces arrived. CENTCOM Adm. Cooper '
    'stated the US has destroyed 17 Iranian vessels since D1.')

add_source_note(doc, ['CTP-ISW Evening Special Report Mar 4', 'CENTCOM'])

add_heading(doc, '1.6  HEZBOLLAH — NORTHERN FRONT RE-OPENED', 3, 'C0392B')

add_body(doc,
    'On 2 March 2026 (D3), Hezbollah Secretary-General Naim Qassem ordered strikes on '
    'Israel following confirmation of Khamenei\'s death, framing the action as "revenge '
    'for the blood of the Supreme Leader." Hezbollah conducted its first strikes on Israel '
    'in more than one year.')

add_bullet(doc, 'Missiles and drone swarms targeted Ramat David Airbase, Meron monitoring base, and Camp Yitzhak.')
add_bullet(doc, 'Israel launched immediate counter-strikes: dozens of airstrikes on Beirut and southern Lebanon.')
add_bullet(doc, 'IDF confirmed assassination of Hezbollah Intelligence Chief Hussein Maklad — responsible for collection architecture and planning targeting of Israel.')
add_bullet(doc, 'Israel struck al-Manar TV and al-Nour Radio stations — Hezbollah\'s information infrastructure.')
add_bullet(doc, 'At least 52 killed and 154 injured in Lebanon from Israeli strikes (March 2–3).', bold_prefix='Lebanese Casualties: ')
add_bullet(doc, '30,000+ displaced in Lebanon; UNHCR reports shelters overwhelmed and civilians sleeping in cars on roadways.')
add_bullet(doc, 'IDF ground forces began operations in southern Lebanon — "forward defence" framing; tens of thousands of reservists mobilizing for potential deeper incursion.')
add_bullet(doc, 'Lebanese government formally banned Hezbollah military activity; President Aoun ordered Hezbollah to hand weapons to state. Hezbollah has not complied.', bold_prefix='Lebanese Government: ')

add_source_note(doc, ['Al Jazeera', 'Times of Israel', 'PBS NewsHour', 'NPR', 'Jerusalem Post', 'Axios', 'CNN'])

add_heading(doc, '1.7  IRAQ PROXIES — ISLAMIC RESISTANCE IN IRAQ', 3, 'C0392B')

add_body(doc,
    'Iran-backed Shiite militias under the Islamic Resistance in Iraq (IRI) umbrella — '
    'including Saraya Awliya al Dam (SAD), Kataib Hezbollah, and Asaib Ahl al Haq — '
    'have significantly escalated attacks on US bases throughout the region.')

add_bullet(doc, 'Erbil International Airport: SAD claimed drone and missile attack on US facilities (March 1). Explosions confirmed.')
add_bullet(doc, 'Camp Victoria (Baghdad IAP): SAD claimed drone attack March 2; air defenses activated.')
add_bullet(doc, 'IRI claimed 28 separate attacks on "enemy bases" on March 2 alone — released footage of launches.')
add_bullet(doc, 'IRI claimed 27 additional missile/drone attacks March 3.')
add_bullet(doc, 'US-Israeli strikes killed 4 PMF personnel at PMF HQ in Diyala Province, Iraq.', bold_prefix='US/Israeli Action: ')
add_bullet(doc, 'CNN correspondent Clarissa Ward: "Iran-backed Iraqi militias are really upping the tempo."')

add_source_note(doc, ['Long War Journal', 'Jerusalem Post', 'CTP-ISW', 'Al Jazeera', 'CNN'])

add_heading(doc, '1.8  HOUTHI — "SOLIDARITY" POSTURE, NOT YET KINETIC', 3, 'C0392B')

add_body(doc,
    'The Houthis (Ansar Allah) occupy a notably restrained posture relative to other '
    'Axis of Resistance members as of D5. CTP-ISW noted: "The Houthis\' continued '
    'relative inaction is notable, given that they were the only Axis of Resistance '
    'member that participated in the June 2025 war."')

add_bullet(doc, 'No confirmed Houthi missile or drone strikes against Israel or shipping as of 04 Mar 06:00 UTC.')
add_bullet(doc, 'Two senior anonymous Houthi officials told AP attacks on Israel and Red Sea shipping will resume (Long War Journal, 28 Feb).')
add_bullet(doc, 'Iranian source told Arab News that Iran "instructed" Houthis to attack Red Sea shipping (28 Feb).')
add_bullet(doc, 'Maersk rerouting vessels via Cape of Good Hope preemptively — Maersk Houston and Astrid Maersk diverted as of March 5 / 12 departures.')
add_bullet(doc, 'UKMTO issued advisory warning elevated risk across Persian Gulf, Gulf of Oman, Strait of Hormuz — including GPS/AIS electronic warfare affecting 1,100+ ships.', bold_prefix='UKMTO: ')
add_bullet(doc, 'BIMCO: war-risk premiums "likely to increase sharply if attacks resume."')

add_callout(doc, '⚠ WATCH:',
    'Houthi activation would open a second maritime strike front simultaneously with Hormuz closure. '
    'Combined effect on Suez and Hormuz chokepoints would constitute near-total Gulf seaborne disruption. '
    'Assess probability of Houthi activation within 72 hours as MODERATE-HIGH.',
    'FEF9E7')

add_source_note(doc, ['Long War Journal', 'CTP-ISW', 'AP via WTOP', 'gCaptain', 'BIMCO', 'UKMTO', 'Maersk'])

doc.add_paragraph()


# ─── D2 NUCLEAR ───────────────────────────────────────────────────────────────

add_domain_header(doc, 'D2', 'NUCLEAR & STRATEGIC PROGRAMS', '8E44AD')

add_heading(doc, '2.1  IAEA VERIFICATION STATUS', 3, '8E44AD')

add_body(doc,
    'IAEA Director General Rafael Grossi has stated unequivocally that the agency '
    '"cannot verify whether Iran has suspended all enrichment-related activities" '
    'or determine "the current size, composition or whereabouts of the stockpile of '
    'enriched uranium in Iran." The IAEA called the "loss of continuity of knowledge" '
    'a matter requiring "utmost urgency."')

add_bullet(doc, 'Iran blocked IAEA inspector access to all sites struck by Israel/US in June 2025, citing a letter (February 2) calling normal safeguards "legally untenable and materially impracticable."')
add_bullet(doc, 'IAEA confirmed Iran did provide access to unaffected nuclear facilities "at least once" since June 2025, except Karun power plant under construction.', bold_prefix='Partial Access: ')
add_bullet(doc, '972 lbs (441 kg) uranium enriched to 60% purity — a short technical step from weapons-grade 90%. IAEA DG Grossi: stockpile sufficient for up to 10 nuclear devices should Iran weaponize.')
add_bullet(doc, 'Most highly enriched uranium stored at underground tunnel complex at Isfahan facility (215 miles SE of Tehran). Isfahan struck in June 2025. IAEA observed "regular vehicular activity" at tunnel complex entrance via commercial satellite.', bold_prefix='Isfahan: ')
add_bullet(doc, 'Natanz: March 3, IAEA confirmed damage to entrance structures renders facility inaccessible. Enrichment status inside unverifiable.', bold_prefix='Natanz: ')
add_bullet(doc, 'Fordow: Activity observed via satellite but "without access... it is not possible to confirm the nature and purpose."', bold_prefix='Fordow: ')

add_source_note(doc, ['IAEA', 'Washington Post', 'PBS NewsHour', 'Al Jazeera', 'ISIS/nuclear-news.net'])

add_heading(doc, '2.2  PRE-WAR NEGOTIATIONS & BREAKDOWN', 3, '8E44AD')

add_body(doc,
    'US-Iran indirect nuclear talks in Geneva (mediated by Oman) concluded two rounds '
    'in February 2026 — Feb 6 and Feb 26. IAEA DG Grossi attended in advisory capacity. '
    'On February 27, IAEA discovered Iran had hidden highly enriched uranium in an '
    'undamaged underground facility from the June 2025 war — a material revelation '
    'that Trump cited in post-strike briefings as justification for military action.')

add_bullet(doc, '"Iran agreed to never, ever have a nuclear material that will create a bomb" (Oman FM, post-talks).')
add_bullet(doc, 'US demanded end to ALL enrichment on Iranian soil plus ballistic missile program dismantlement — red lines Iran could not accept publicly.')
add_bullet(doc, 'February 13: Trump said regime change "would be the best thing that could happen."')
add_bullet(doc, 'February 14: US military briefed reporters on "weeks-long sustained operations" planning — signaling strikes were already decided.')

add_source_note(doc, ['Wikipedia 2025-2026 negotiations', 'Al Jazeera', 'IAEA', 'WaPo'])

add_heading(doc, '2.3  IRANIAN NUCLEAR PROGRAM RESILIENCE', 3, '8E44AD')

add_bullet(doc, 'Iran deliberately dispersed centrifuge components across multiple production sites and embedded manufacturing in hardened tunnels. Even comprehensive strike campaigns cannot achieve certainty of destruction.')
add_bullet(doc, 'HEU stockpile location remains UNKNOWN to IAEA — and therefore uncertain to strike planners. If relocated prior to strikes, a reconstitutable program remains possible.')
add_bullet(doc, 'Iran\'s baseline position: "We are not pursuing nuclear weapons." Has resisted all demands to halt enrichment or transfer HEU stockpile out of country.')

add_callout(doc, '◾ ASSESSMENT:',
    'Iran has not declared a nuclear weapon. However, the combination of 60% HEU stockpile, '
    'blocked inspections, and active weaponization research strikes creates genuine breakout risk '
    'that cannot be assessed without restored IAEA access. The military campaign has degraded '
    'known nuclear infrastructure but has not eliminated the underlying enrichment knowledge or '
    'all fissile material. Nuclear verification remains the single most important outstanding unknown.',
    'EAD7F2')

add_source_note(doc, ['IAEA', 'ISIS Online', 'Washington Post', 'Al Jazeera'])

doc.add_paragraph()


# ─── D3 ENERGY ────────────────────────────────────────────────────────────────

add_domain_header(doc, 'D3', 'ENERGY & ECONOMIC DISRUPTION', '1A5276')

add_heading(doc, '3.1  STRAIT OF HORMUZ — DE FACTO CLOSURE', 3, '1A5276')

add_body(doc,
    'The Strait of Hormuz, through which approximately 20% of global seaborne oil '
    'and 12–14 million bbl/day of crude and condensate transit, is effectively closed '
    'to commercial shipping. IRGC confirmed formal closure March 2; insurance withdrawal '
    'has achieved de facto blockade without requiring a sustained naval force.')

add_bullet(doc, 'Tanker transit dropped ~70%, then to near-zero; 150+ vessels anchored outside the strait awaiting resolution.')
add_bullet(doc, 'No major carrier has transited since March 1. Maersk, MSC, CMA-CGM rerouting via Cape of Good Hope.')
add_bullet(doc, 'War-risk insurance premiums spiked from ~0.2% to 2%+ of vessel value per voyage — making transit financially untenable for most operators.', bold_prefix='Insurance: ')
add_bullet(doc, 'Electronic warfare: GPS/AIS disruption affecting 1,100+ ships across Gulf, Gulf of Oman, and Hormuz (UKMTO).')

add_source_note(doc, ['CNBC', 'Al Jazeera', 'Kpler Mar 1', 'Wood Mackenzie', 'gCaptain', 'UKMTO', 'Wikipedia 2026 Hormuz crisis'])

add_heading(doc, '3.2  OIL & ENERGY PRICE MOVEMENTS', 3, '1A5276')

add_body(doc, 'Crude prices have undergone the sharpest weekly rise since 2022 Russia-Ukraine invasion.')

facts_table2 = doc.add_table(rows=6, cols=3)
facts_table2.style = 'Table Grid'
facts_table2.autofit = False

for i, (label, val, note) in enumerate([
    ('Commodity', 'Price (04 Mar)', 'Change'),
    ('Brent Crude', '$83.83/bbl', '+25% from pre-war level'),
    ('WTI Crude', '$76.59/bbl', '+8.4% on Day 1 alone'),
    ('European Natural Gas', 'Surged 20%+', 'LNG transit threatened'),
    ('US Gasoline (avg)', '+10–30¢/gallon (est.)', 'GasBuddy est.; some markets +85¢'),
    ('Henry Hub Gas', '$7.72/MMBtu', '+81% vs. December 2025'),
]):
    row = facts_table2.rows[i]
    if i == 0:
        for j in range(3):
            set_cell_bg(row.cells[j], '2C3E50')
            p = row.cells[j].paragraphs[0]
            r = p.add_run(val if j == 1 else (label if j == 0 else note))
            r.bold = True; r.font.size = Pt(9); r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    else:
        for j, text in enumerate([label, val, note]):
            p = row.cells[j].paragraphs[0]
            r = p.add_run(text)
            r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
            if j == 0: r.bold = True

doc.add_paragraph()

add_bullet(doc, 'Barclays: Brent could hit $100/bbl as security situation spirals.', bold_prefix='Analyst Forecasts: ')
add_bullet(doc, 'UBS: Material disruption scenario could send Brent above $120/bbl.')
add_bullet(doc, 'Wood Mackenzie (Mar 2 emergency webinar): "$150 oil possible if Hormuz remains shut." 12–14 million bbl/day of crude and condensate + 30% of Europe\'s jet fuel + 20% of global LNG supply at risk.')
add_bullet(doc, 'OPEC+ pledged +206,000 bbl/day emergency output increase — analysts note this is insufficient to compensate for Hormuz closure.')
add_bullet(doc, 'Russia competitive position improving materially — India and China face incentives to deepen reliance on Russian supply.', bold_prefix='Secondary Effect: ')
add_bullet(doc, 'Pakistan, Bangladesh, India most vulnerable LNG import markets — Qatar and UAE supply ~99%, 72%, 53% of their LNG respectively. Limited storage and procurement alternatives.', bold_prefix='LNG Vulnerability: ')

add_source_note(doc, ['CNBC', 'NPR', 'Al Jazeera Economy', 'Kpler', 'Wood Mackenzie', 'Barclays', 'UBS', 'EIA', 'Goldman Sachs'])

add_heading(doc, '3.3  REGIONAL ENERGY INFRASTRUCTURE STRUCK', 3, '1A5276')

add_bullet(doc, 'Jebel Ali Port, UAE: debris fires from intercepted missiles; temporary disruption of a primary Gulf commercial hub.')
add_bullet(doc, 'Duqm Port, Oman: Iranian drone struck fuel tanker.')
add_bullet(doc, 'Port of Salalah, Oman: Iranian drone strike on port facility.')
add_bullet(doc, 'Kuwait International Airport: struck; air defenses engaged.')
add_bullet(doc, 'Dubai International Airport: struck; evacuated; "minor damage" per UAE authorities; 4 staff injured.')
add_bullet(doc, 'AWS/Amazon data center, UAE (mec1-az2): direct strike, fire, power shutdown — cloud infrastructure disruption for ME region.')

add_source_note(doc, ['Euronews', 'CNBC', 'CENTCOM', 'Breaking Defense'])

add_heading(doc, '3.4  TRUMP TANKER ESCORT PROPOSAL', 3, '1A5276')

add_body(doc,
    'President Trump stated: "If necessary, the United States Navy will begin escorting '
    'tankers through the Strait of Hormuz, as soon as possible." No operational order '
    'has been announced. The US Fifth Fleet has pre-positioned vessels; however, '
    'escorting commercial tankers through an active war zone under IRGC missile and '
    'drone threat would expose warships to significant attrition risk without resolving '
    'the insurance-withdrawal constraint.')

add_source_note(doc, ['CNBC', 'Al Jazeera', 'NPR'])

doc.add_paragraph()


# ─── D4 DIPLOMATIC ────────────────────────────────────────────────────────────

add_domain_header(doc, 'D4', 'DIPLOMATIC & POLITICAL', '145A32')

add_heading(doc, '4.1  US DIPLOMATIC POSTURE', 3, '145A32')

add_bullet(doc, 'Trump March 3: "It\'s too late for negotiations with Iran" — then within 24 hours, "They want to talk, and I have agreed to talk."', bold_prefix='Contradictory Signals: ')
add_bullet(doc, 'Trump stated objective: "Destroy Iran\'s missile and military capabilities, prevent Iran from obtaining nuclear weapons, and achieve regime change."')
add_bullet(doc, 'Timeline stated: "Likely one month or less." Trump has "not ruled out" use of ground forces.')
add_bullet(doc, 'Trump offered new rationale: Iran had hidden HEU from inspectors; this was presented as justification for pre-empting nuclear breakout.')
add_bullet(doc, 'State Department: ordered non-emergency government personnel and families to depart Jordan, Bahrain, Iraq, Qatar, Kuwait, and UAE.', bold_prefix='Evacuation Orders: ')

add_source_note(doc, ['NPR', 'CBS News', 'PBS NewsHour', 'CNN', 'Al Jazeera', 'Pravda EN (Trump/Iran quotes)'])

add_heading(doc, '4.2  IRAN\'S LEADERSHIP & SUCCESSION', 3, '145A32')

add_body(doc,
    'Supreme Leader Ali Khamenei was killed D2 (confirmed 1 March). His wife died of '
    'injuries sustained in the strike. A provisional Leadership Council has assumed '
    'authority. The Assembly of Experts — the body constitutionally tasked with electing '
    'a new supreme leader — was itself struck during an emergency session, disrupting '
    'succession proceedings.')

add_bullet(doc, 'Iran\'s Security Chief Ali Larijani: "We will not negotiate with the United States." Contradicts Iran\'s new leadership\'s private approach to Trump.')
add_bullet(doc, 'Reports: Iran\'s new leadership internally sought to resume negotiations.')
add_bullet(doc, 'Larijani\'s public rejection vs. private overtures signals internal disagreement in interim leadership.')

add_source_note(doc, ['Wikipedia', 'The Atlantic via Pravda EN', 'CNBC', 'CTP-ISW'])

add_heading(doc, '4.3  UN SECURITY COUNCIL', 3, '145A32')

add_bullet(doc, 'Emergency session convened February 28 at request of France, Colombia, China, Russia.', bold_prefix='Emergency Session: ')
add_bullet(doc, 'UN Secretary-General Guterres warned of "a chain of events that nobody can control" — called strikes a "squandering" of diplomatic opportunity.')
add_bullet(doc, 'No ceasefire resolution tabled — US veto is near-certain on any resolution condemning its own operations. US holds Security Council presidency in March 2026.')
add_bullet(doc, 'DAWN advocacy called for UN General Assembly emergency session under "Uniting for Peace" procedure — not yet initiated.')
add_bullet(doc, 'UNSC Resolution 2812 (2026) extended Houthi Red Sea reporting mandate for 6 months (pre-war, passed earlier in March).', bold_prefix='Related Action: ')

add_source_note(doc, ['UN News', 'UN Press', 'DAWN', 'Security Council Report', 'House of Commons Library'])

add_heading(doc, '4.4  INTERNATIONAL REACTIONS', 3, '145A32')

add_bullet(doc, 'France 24 / EU: Commission President von der Leyen appeared to support Trump\'s regime-change push, calling for a "credible transition."', bold_prefix='EU: ')
add_bullet(doc, 'Emmanuel Macron: called for urgent UNSC meeting; "serious consequences" for international peace.', bold_prefix='France: ')
add_bullet(doc, '"Immediate ceasefire" demand; condemned killing of Khamenei as violation of UN Charter.', bold_prefix='China: ')
add_bullet(doc, 'Putin expressed condolences; foreign ministry called for "immediate return to political and diplomatic track."', bold_prefix='Russia: ')
add_bullet(doc, 'UK PM Starmer: US cannot use British airbases (Diego Garcia, RAF Fairford) for Iran strikes — would breach international law.', bold_prefix='UK: ')
add_bullet(doc, 'UAE, Qatar, Saudi Arabia, Kuwait, Bahrain, Jordan: joint statement condemning Iran\'s "indiscriminate and reckless" strikes on their sovereign territory; pledged to "stand united in defense."', bold_prefix='Gulf States: ')
add_bullet(doc, 'UAE considering military action against Iranian missile sites — "unprecedented" and reflecting "enormous anger."', bold_prefix='UAE Escalation Risk: ')

add_source_note(doc, ['CNN', 'Al Jazeera', 'CNBC', 'The Guardian', 'BBC', 'UAE MOFA', 'Axios', 'France 24', 'TIME'])

doc.add_paragraph()


# ─── D5 CYBER ─────────────────────────────────────────────────────────────────

add_domain_header(doc, 'D5', 'CYBER, INFORMATION OPS & ELECTRONIC WARFARE', '1F618D')

add_heading(doc, '5.1  IRAN INTERNET BLACKOUT', 3, '1F618D')

add_body(doc,
    'Iran has undergone one of the most severe internet shutdowns ever recorded. '
    'NetBlocks confirmed connectivity at approximately 1% of normal levels as of '
    'D4 (March 3–4). This constitutes a regime-imposed "National Information Network" '
    'isolation — not damage from strikes.')

add_bullet(doc, 'Connectivity dropped from ~50% at start of strikes (Feb 28) to ~1% within hours. Georgia Tech Internet Intelligence Lab: "most sophisticated and severe in Iran\'s history."')
add_bullet(doc, 'Economic cost: NetBlocks estimates $37M+ daily loss. Online sales -80%. Tehran Stock Exchange lost 450,000 points in 4 days.', bold_prefix='Economic Cost: ')
add_bullet(doc, '"Whitelisting" system maintains access for regime-loyal entities and essential state functions only.', bold_prefix='Architecture: ')
add_bullet(doc, '"Barracks Internet" long-term plan: access permanently limited to approved entities. Government spokesperson: international internet will "never return to its previous form."', bold_prefix='Permanent Change: ')
add_bullet(doc, 'US reportedly pre-deployed 7,000+ Starlink terminals into Iran prior to war (Wall Street Journal). Iranian authorities moving to jam/seize terminals.', bold_prefix='Starlink: ')

add_source_note(doc, ['NetBlocks', 'CNBC', 'The Register', 'Rest of World', 'The National'])

add_heading(doc, '5.2  ISRAEL\'S CYBER CAMPAIGN', 3, '1F618D')

add_body(doc,
    'Parallel to kinetic strikes, Israel launched what analysts described as "the largest '
    'cyberattack in history" against Iran. Targeting included: government communication '
    'systems, state media, energy/aviation infrastructure, and IRGC command networks. '
    'This directly contributed to the internet blackout and media blackout affecting '
    'Iranian population.')

add_source_note(doc, ['Cybernews', 'SC Media', 'CSIS', 'CTP-ISW'])

add_heading(doc, '5.3  IRANIAN CYBER THREAT — ELEVATED STATE', 3, '1F618D')

add_body(doc,
    'Multiple intelligence and threat research organizations have elevated Iran-linked '
    'cyber threat levels to highest-alert status. The "Electronic Operations Room" — '
    'established February 28, 2026 — coordinates IRGC-aligned APTs and hacktivist proxies.')

add_bullet(doc, 'APT33, APT34 (OilRig), APT35 (Charming Kitten), APT42, MuddyWater — all identified in elevated reconnaissance and espionage posture.', bold_prefix='State APTs: ')
add_bullet(doc, 'Cotton Sandstorm (IRGC-linked): deploying WezRat infostealer via spearphishing masquerading as urgent software updates (Check Point).', bold_prefix='WezRat Deployment: ')
add_bullet(doc, 'Handala Hack (MOIS-linked hacktivist): data exfiltration and cyber disruption operations.', bold_prefix='Hacktivist Proxies: ')
add_bullet(doc, 'Hydro Kitten: "specific threats targeting financial services sector" (security researchers).', bold_prefix='Financial Sector: ')
add_bullet(doc, 'DHS bulletin: Iran-aligned hacktivists could conduct low-level attacks on US networks; large-scale physical attack unlikely.', bold_prefix='DHS Warning: ')
add_bullet(doc, 'NCSC (UK): issued advisory urging UK organisations to take action following Middle East escalation.', bold_prefix='NCSC: ')
add_bullet(doc, 'Palo Alto Unit 42: published "Threat Brief: March 2026 Escalation of Cyber Risk Related to Iran" — rapid mitigation of external vulnerabilities advised.', bold_prefix='Unit 42: ')

add_source_note(doc, ['CISA', 'Unit 42 Palo Alto', 'NCSC', 'The Register', 'Cybernews', 'SC Media', 'Cybersecurity Dive', 'CloudSEK', 'CSIS', 'Check Point'])

add_heading(doc, '5.4  ELECTRONIC WARFARE', 3, '1F618D')

add_bullet(doc, 'GPS and AIS spoofing/jamming affecting 1,100+ vessels across Persian Gulf, Gulf of Oman, Strait of Hormuz, UAE, Qatari, and Omani waters (UKMTO advisory).', bold_prefix='Maritime EW: ')
add_bullet(doc, 'Pattern: sophisticated EW probing of regional APIs and government infrastructure began early February, stopped February 27 — likely linked to onset of internet blackout.', bold_prefix='Pre-Conflict Probing: ')
add_bullet(doc, 'Recommended defensive posture (CISA/Unit 42): immediately mitigate external-facing vulnerabilities; ensure OT/ICS systems not connected to public internet; assume phishing and supply-chain vectors active.', bold_prefix='Defensive Actions: ')

add_source_note(doc, ['UKMTO', 'CISA', 'Unit 42', 'CloudSEK', 'Cybernews'])

doc.add_paragraph()


# ─── OUTLOOK / ASSESSMENT ─────────────────────────────────────────────────────

add_heading(doc, '◾ ASSESSMENT — 72-HOUR OUTLOOK', 2, '1A1A2E')
add_horizontal_rule(doc)

add_body(doc,
    'The conflict is in a critical phase. Iran\'s offensive ballistic missile launch '
    'tempo is visibly declining, which may reflect launcher attrition — a positive '
    'indicator for the campaign\'s stated objective of degrading Iran\'s ballistic '
    'missile capacity. Simultaneously, US interceptor stockpiles face genuine strain. '
    'The resolution of the "missile math" problem — who runs out first — is the '
    'primary tactical variable in the next 72 hours.', bold=False)

add_bullet(doc, 'HIGH — Iran will attempt to sustain proxy pressure via Iraqi militias and Hezbollah as its direct strike tempo slows.', bold_prefix='Proxy Escalation: ')
add_bullet(doc, 'MODERATE-HIGH — Houthi maritime activation within 72 hours; could compound Hormuz closure with Bab el-Mandeb threat.', bold_prefix='Houthi Activation: ')
add_bullet(doc, 'MODERATE — UAE military response to Iranian strikes; unprecedented but Emirati rhetoric is escalatory.', bold_prefix='UAE Military Action: ')
add_bullet(doc, 'MODERATE — US-Iran ceasefire/negotiation opening; contradictory signals from both sides suggest back-channel contact is ongoing.', bold_prefix='Diplomatic Channel: ')
add_bullet(doc, 'HIGH — IRGC/APT cyber disruption targeting US financial services, energy, and water sectors.', bold_prefix='Cyber Strike on US: ')
add_bullet(doc, 'LOW-MODERATE — Iranian nuclear breakout attempt; HEU stockpile location unknown; if Iran retains protected centrifuge capacity, this risk persists but is not the immediate focus.', bold_prefix='Nuclear: ')
add_bullet(doc, 'CERTAIN — Brent crude above $85/bbl by 5 March absent ceasefire announcement; $100+ threshold possible if Hormuz closure extends 14+ days.', bold_prefix='Energy Markets: ')

doc.add_paragraph()


# ─── SOURCE ATTRIBUTION ───────────────────────────────────────────────────────

add_heading(doc, '◾ SOURCE ATTRIBUTION', 2, '1A1A2E')
add_horizontal_rule(doc)

add_body(doc,
    'This brief was compiled from the following sources, all consulted on 4 March 2026. '
    'Where information was reported by a single source, it is hedged accordingly in-text.', italic=True)

source_table = doc.add_table(rows=1, cols=2)
source_table.style = 'Table Grid'
source_table.autofit = False
source_table.columns[0].width = Inches(3.3)
source_table.columns[1].width = Inches(3.3)

set_cell_bg(source_table.rows[0].cells[0], '2C3E50')
set_cell_bg(source_table.rows[0].cells[1], '2C3E50')
for j, header in enumerate(['TIER 1 / PRIMARY', 'TIER 2–3 / ANALYTICAL']):
    p = source_table.rows[0].cells[j].paragraphs[0]
    r = p.add_run(header)
    r.bold = True; r.font.size = Pt(9); r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

tier1_row = source_table.add_row()
tier2_row = source_table.add_row()

for cell, text in [
    (tier1_row.cells[0],
     'Associated Press (AP)\nReuters Middle East\nCTP-ISW (Critical Threats Project / ISW)\nIAEA\nCENTCOM\nUKMTO\nIDF Spokesperson\nAl Jazeera English\nLong War Journal\nNetBlocks'),
    (tier1_row.cells[1],
     'CNN · BBC · CNBC · NYT · The Guardian\nTimes of Israel · Jerusalem Post\nMiddle East Eye · France 24 · DW\nNPR · PBS NewsHour · ABC News\nThe Register · SC Media · CSIS · CSIS\nUnit 42 (Palo Alto Networks) · NCSC\nCybernews · Cybersecurity Dive\ngCaptain · BIMCO · Kpler · Wood Mackenzie\nBarclays / UBS / Goldman Sachs (analyst notes)\nTime · Axios · Breaking Defense')
]:
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)

doc.add_paragraph()


# ─── FOOTER ───────────────────────────────────────────────────────────────────

add_horizontal_rule(doc)
footer_p = doc.add_paragraph()
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_run = footer_p.add_run(
    'CSE Intel Brief  ·  4 March 2026  ·  Morning Edition  ·  '
    'UNCLASSIFIED // FOR OFFICIAL USE ONLY\n'
    'Not for public release. All information from open-source and publicly available reporting. '
    'All casualty figures and operational claims should be independently verified before action.'
)
footer_run.italic = True
footer_run.font.size = Pt(8)
footer_run.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)


# ─── SAVE ─────────────────────────────────────────────────────────────────────

output_path = 'briefs/CSE_Intel_Brief_20260304.docx'
os.makedirs('briefs', exist_ok=True)
doc.save(output_path)
print(f'✓ Brief saved: {output_path}')
